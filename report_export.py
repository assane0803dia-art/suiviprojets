"""
Convertit un rapport texte (format Markdown simple : ## titres, - listes, **gras**)
en fichiers Word (.docx) et PDF téléchargeables — avec une vraie mise en forme
(le **gras** Markdown devient du gras réel, pas des astérisques littéraux).
"""

import io
import re
from xml.sax.saxutils import escape
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm

BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")


def _parse_lines(report_text):
    """Découpe le rapport en lignes typées : (type, texte) où type est 'h2', 'bullet' ou 'p'."""
    parsed = []
    for raw_line in report_text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("## "):
            parsed.append(("h2", line[3:].strip()))
        elif line.startswith("# "):
            parsed.append(("h1", line[2:].strip()))
        elif line.startswith("- ") or line.startswith("• "):
            parsed.append(("bullet", line[2:].strip()))
        else:
            parsed.append(("p", line))
    return parsed


def _add_runs_docx(paragraph, text):
    """Ajoute le texte à un paragraphe Word en convertissant **gras** en runs réellement en gras."""
    pos = 0
    for m in BOLD_PATTERN.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _strip_bold_markers(text):
    """Pour les titres : retire les ** sans les convertir (les titres sont déjà en gras par leur style)."""
    return BOLD_PATTERN.sub(r"\1", text)


def _to_reportlab_markup(text):
    """Échappe le texte puis convertit **gras** en balises <b> reconnues par reportlab."""
    escaped = escape(text)
    return BOLD_PATTERN.sub(r"<b>\1</b>", escaped)


def export_to_docx(report_text: str, projet_nom: str) -> bytes:
    doc = Document()
    doc.add_heading(f"Rapport d'exécution — {projet_nom}", level=0)

    for kind, text in _parse_lines(report_text):
        if kind == "h1":
            doc.add_heading(_strip_bold_markers(text), level=1)
        elif kind == "h2":
            doc.add_heading(_strip_bold_markers(text), level=2)
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_docx(p, text)
        else:
            p = doc.add_paragraph()
            _add_runs_docx(p, text)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_to_pdf(report_text: str, projet_nom: str) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleCustom", parent=styles["Title"], fontSize=16, spaceAfter=16)
    h2_style = ParagraphStyle("H2Custom", parent=styles["Heading2"], spaceBefore=12, spaceAfter=6)
    bullet_style = ParagraphStyle("BulletCustom", parent=styles["Normal"], leftIndent=14, spaceAfter=4)
    normal_style = ParagraphStyle("NormalCustom", parent=styles["Normal"], spaceAfter=6)

    story = [Paragraph(escape(f"Rapport d'exécution — {projet_nom}"), title_style), Spacer(1, 6)]

    for kind, text in _parse_lines(report_text):
        if kind in ("h1", "h2"):
            story.append(Paragraph(escape(_strip_bold_markers(text)), h2_style))
        elif kind == "bullet":
            story.append(Paragraph(f"• {_to_reportlab_markup(text)}", bullet_style))
        else:
            story.append(Paragraph(_to_reportlab_markup(text), normal_style))

    doc.build(story)
    return buffer.getvalue()
